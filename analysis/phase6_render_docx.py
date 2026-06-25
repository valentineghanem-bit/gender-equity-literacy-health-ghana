"""
Render manuscript_draft_v3.md -> Q1 .docx with EMBEDDED TABLES (real Word tables) and FIGURES.
Times New Roman 12pt, 1.5 spacing, justified; tables from outputs/tables/manuscript_table{1,2}.csv;
figures from outputs/figures; references appended from v1. Local only (Tenet 20).
"""
import re
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
M = ROOT/"manuscript"; F = ROOT/"outputs"/"figures"; T = ROOT/"outputs"/"tables"
SRC = M/"manuscript_draft_v3.md"; V1 = M/"manuscript_draft_v1.md"
OUT = M/"Gender_Equity_Literacy_Health_Ghana_v3.docx"

FIGMAP = {"1":["fig_choropleth_illiteracy_261.png","fig_choropleth_poverty_261.png"],
          "2":["fig_lisa_illiteracy_261.png","fig_lisa_poverty_261.png"],
          "3":["shap_summary_district.png"], "4":["fig04_multivariate_disparity.png"],
          "5":["fig_mediation_forest.png"]}
TABMAP = {"1":"manuscript_table1.csv", "2":"manuscript_table2.csv"}

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
st.paragraph_format.line_spacing = 1.5; st.paragraph_format.space_after = Pt(6)

def runs(p, text):
    for i, seg in enumerate(re.split(r"\*\*", text)):
        if seg: r = p.add_run(seg); r.bold = (i % 2 == 1)

def heading(text, size, before=10):
    p = doc.add_paragraph(); r = p.add_run(re.sub(r"\*\*","",text)); r.bold=True; r.font.size=Pt(size)
    p.paragraph_format.space_before = Pt(before)

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # bold the leading "Table N." / "Figure N."
    m = re.match(r"(Table \d+\.|Figure \d+\.)\s*(.*)", text)
    if m:
        r=p.add_run(m.group(1)+" "); r.bold=True; r.font.size=Pt(10)
        r2=p.add_run(m.group(2)); r2.italic=True; r2.font.size=Pt(10)
    else:
        r=p.add_run(text); r.italic=True; r.font.size=Pt(10)

def add_table_csv(csv):
    df = pd.read_csv(T/csv).fillna("")
    tb = doc.add_table(rows=1, cols=len(df.columns)); tb.style = "Table Grid"
    for j,c in enumerate(df.columns):
        cell = tb.rows[0].cells[j]; run = cell.paragraphs[0].add_run(str(c))
        run.bold=True; run.font.size=Pt(8.5)
    for _,row in df.iterrows():
        cells = tb.add_row().cells
        for j,v in enumerate(row):
            r = cells[j].paragraphs[0].add_run(str(v)); r.font.size=Pt(8.5)

def embed_fig(num):
    for png in FIGMAP.get(num, []):
        fp = F/png
        if fp.exists():
            doc.add_picture(str(fp), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

mode = None
for ln in SRC.read_text(encoding="utf-8").splitlines():
    s = ln.rstrip()
    if not s.strip() or s.strip()=="---": continue
    if s.startswith("**Draft v3"): continue
    if s.startswith("## References"): break
    if s.startswith("## Tables"): mode="tables"; heading("Tables",13); continue
    if s.startswith("## Figure legends"): mode="figs"; heading("Figures",13); continue
    if mode=="tables" and s.startswith("- **Table"):
        m=re.match(r"- \*\*Table (\d+)\.\*\*\s*(.*)", s)
        if m:
            caption("Table %s. %s"%(m.group(1), re.sub(r"\*\*","",m.group(2))))
            add_table_csv(TABMAP[m.group(1)]); doc.add_paragraph(); continue
    if mode=="figs" and s.startswith("- **Figure"):
        m=re.match(r"- \*\*Figure (\d+)\.\*\*\s*(.*)", s)
        if m: embed_fig(m.group(1)); caption("Figure %s. %s"%(m.group(1), re.sub(r"\*\*","",m.group(2)))); continue
    if s.startswith("# "):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(s[2:]); r.bold=True; r.font.size=Pt(15)
    elif s.startswith("## "): heading(s[3:],13)
    elif s.startswith("### "): heading(s[4:],12,4)
    elif s.startswith("- "):
        p=doc.add_paragraph(style="List Bullet"); runs(p,s[2:])
    else:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; runs(p,s)

# references from v1
heading("References",13); grab=False
for ln in V1.read_text(encoding="utf-8").splitlines():
    s=ln.rstrip()
    if s.startswith("## References"): grab=True; continue
    if grab:
        if s.startswith("*Data sources:*"):
            p=doc.add_paragraph(); r=p.add_run(re.sub(r"[*]","",s)); r.italic=True; r.font.size=Pt(10); break
        if s.startswith("## "): break
        if re.match(r"^\d+\.\s", s):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Pt(18); p.paragraph_format.first_line_indent=Pt(-18)
            p.add_run(s).font.size=Pt(11)

f=doc.sections[0].footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); f._p.append(fld)
doc.save(str(OUT))
import os
print(f"[docx] {OUT.name} ({round(os.path.getsize(OUT)/1024)} KB) | images={len(doc.inline_shapes)} | tables={len(doc.tables)}")
